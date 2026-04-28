import os
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from pathlib import Path

def fix_template():
    template_path = Path("docs/writing/templates/Plantilla presentación trabajos.docx")
    output_path = Path("docs/writing/templates/Master_Septimo_Ciclo.docx")
    
    if not template_path.exists():
        print(f"Error: No se encuentra la plantilla en {template_path}")
        return

    doc = Document(template_path)
    
    replacements = {
        "Quinto A": "Séptimo Ciclo",
        "Nombre: (colocar su nombre completo en caso de trabajo individual, o los nombres de los integrantes, en orden alfabético, separados con coma (,))": "Nombre: Erick Fabricio Condoy Seraquive",
        "(fecha de entrega)": "24 de Abril de 2026",
        "Septiembre 2024-Febrero 2025": "Marzo 2026 - Agosto 2026"
    }

    # Procesar párrafos
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # Eliminar resaltado amarillo
            if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                run.font.highlight_color = None
            
            # Reemplazar texto
            for old_text, new_text in replacements.items():
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)

    # También procesar tablas (a veces el encabezado está en tablas)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                            run.font.highlight_color = None
                        for old_text, new_text in replacements.items():
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)

    doc.save(output_path)
    print(f"Plantilla curada guardada en: {output_path}")

if __name__ == "__main__":
    fix_template()

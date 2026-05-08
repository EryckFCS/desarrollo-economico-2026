# /// script
# dependencies = [
#   "python-docx",
# ]
# ///

import os
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def replace_text_in_paragraph(paragraph, old_text, new_text):
    """Surgical replacement in runs to preserve formatting."""
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)

def fill_report():
    template_path = "Taller 2_U1_plantilla presentación de informe (1).docx"
    output_path = "Taller 2_U1_Informe_Final_Grupo3.docx"
    
    if not os.path.exists(template_path):
        print(f"Error: Template {template_path} not found.")
        return

    doc = Document(template_path)

    # 1. Configuración de Márgenes (APA: 2.54 cm)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # 2. Reemplazo Quirúrgico de Identificación (Preservando Negritas/Estilo)
    for p in doc.paragraphs:
        replace_text_in_paragraph(p, "(Nombre de la teoría asignada)", "Teoría Neomarxista de la Dependencia")
        replace_text_in_paragraph(p, "Nombre:", "Nombre: Grupo 3 (Séptimo A)")
        replace_text_in_paragraph(p, "Fecha:", "Fecha: 04 de mayo de 2026")
        replace_text_in_paragraph(p, "(Proponer un título creativo)", "Crudo Castigo: La Anatomía Política del Valor")
        replace_text_in_paragraph(p, "(Describir brevemente dónde y cuándo se desarrolla la historia)", "Venezuela, 1943. Un tribunal transhistórico audita la estructura de dependencia.")

    # 3. Inserción de Contenido Extenso con Sangría
    # Buscamos los párrafos que contienen las instrucciones para reemplazarlos
    for p in doc.paragraphs:
        # Definición
        if "(Explicar en un párrafo breve" in p.text:
            p.text = "" # Borramos el texto de instrucción pero mantenemos el párrafo
            run = p.add_run("La teoría neomarxista del desarrollo sostiene que el subdesarrollo no es un estado transitorio, sino un proceso activo generado por la expansión del capitalismo industrial. El sistema mundial funciona como un mecanismo de succión donde la acumulación de los países centrales se sustenta en la explotación de la periferia (Reyes, 2001).")
            p.paragraph_format.first_line_indent = Cm(1.27)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Características
        if "Característica 1:" in p.text:
            p.text = ""
            p.add_run("- Desarrollo del Subdesarrollo: André Gunder Frank (1967) postula que el subdesarrollo fue creado por el mismo proceso histórico que generó el desarrollo del centro.")
            p.paragraph_format.first_line_indent = Cm(1.27)
        if "Característica 2:" in p.text:
            p.text = ""
            p.add_run("- Intercambio Desigual: Samir Amin (1976) explica la transferencia neta de valor desde la periferia hacia el centro debido a asimetrías tecnológicas.")
            p.paragraph_format.first_line_indent = Cm(1.27)
        if "Característica 3:" in p.text:
            p.text = ""
            p.add_run("- Alianzas de Clase: La dependencia se reproduce mediante alianzas entre el capital transnacional y las élites locales (lumpenburguesía).")
            p.paragraph_format.first_line_indent = Cm(1.27)
        
        # Limpieza de instrucciones
        if "(Pueden incluir de 3 a 5" in p.text or "(Se sugiere indicar entre 2 y 4" in p.text:
            p.text = ""

        # Políticas
        if "Política 1:" in p.text:
            p.text = ""
            p.add_run("- La Desconexión (Delinking): Subordinar las relaciones externas a las necesidades del desarrollo interno (Amin, 1976).")
            p.paragraph_format.first_line_indent = Cm(1.27)
        if "Política 2:" in p.text:
            p.text = ""
            p.add_run("- Nacionalización Estratégica: Control estatal de sectores clave para retener el excedente económico.")
            p.paragraph_format.first_line_indent = Cm(1.27)
        if "Política 3:" in p.text:
            p.text = ""
            p.add_run("- Cooperación Sur-Sur: Alianzas con otros países periféricos para romper el monopolio tecnológico.")
            p.paragraph_format.first_line_indent = Cm(1.27)

    # 4. Reparto de Personajes
    # Buscamos el párrafo "Personajes:" para insertar después
    target_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if "Personajes:" in p.text:
            target_idx = i
            break
    
    if target_idx != -1:
        # Borramos las balas de ejemplo
        for j in range(target_idx + 1, target_idx + 6):
            if j < len(doc.paragraphs) and ("(Nombre del personaje" in doc.paragraphs[j].text or doc.paragraphs[j].text.strip() == "-"):
                doc.paragraphs[j].text = ""
        
        characters = [
            "Su Señoría (Jefferson Vasquez): Juez de Sistemas.",
            "Fiscal de la Soberanía (Erick Condoy): Analista de asimetrías.",
            "Secretaria de Actas (Kristie Japon): Auditora del sistema.",
            "Sr. Williams (Erik Moreno): Representante del Centro (Optimización de Capital).",
            "Dr. Mendoza (Luis Hurtado): Intermediario local.",
            "Juan (Luis Abad): Variable de ajuste (Trabajo).",
            "María (Gabriela Palacios): Recurso base (Naturaleza)."
        ]
        # Insertamos después de "Personajes:"
        for char in reversed(characters):
            new_p = doc.paragraphs[target_idx].insert_paragraph_before(f"- {char}")
            new_p.style = doc.styles['List Paragraph']
            new_p.paragraph_format.first_line_indent = Cm(1.27)

    # 5. Escenas y Conclusión
    scenes = {
        "Escena 1:": "LA LÓGICA DEL CAPITAL. El Fiscal acusa al Dr. Mendoza de entregar el subsuelo mediante la Ley de 1943.",
        "Escena 2:": "LA VARIABLE DE AJUSTE. Testimonios de Juan (Obrero) y María (Tierra) sobre la dependencia técnica.",
        "Escena 3:": "EL VEREDICTO. El Tribunal ordena la 'Desconexión Soberana' frente a las amenazas del Sr. Williams."
    }
    
    for p in doc.paragraphs:
        for s_key, s_desc in scenes.items():
            if s_key in p.text:
                p.text = ""
                run = p.add_run(f"{s_key} {s_desc}")
                run.bold = True
                p.paragraph_format.first_line_indent = Cm(1.27)
        
        if "Conclusión o mensaje final:" in p.text:
            p.text = "Conclusión: El desarrollo requiere soberanía tecnológica. 'El precio de la luz propia es atravesar la sombra del Centro'."
            p.paragraph_format.first_line_indent = Cm(1.27)
        
        if "Bibliografía:" in p.text:
            # Limpiamos y añadimos
            bib_text = (
                "Amin, S. (1976). Unequal Development. Monthly Review Press.\n"
                "Frank, A. G. (1967). Capitalism and Underdevelopment. Monthly Review Press.\n"
                "Reyes, G. E. (2001). Principales teorías sobre el desarrollo económico. Nómadas, (4)."
            )
            doc.add_paragraph(bib_text)

    doc.save(output_path)
    print(f"Report optimized and generated: {output_path}")

if __name__ == "__main__":
    fill_report()

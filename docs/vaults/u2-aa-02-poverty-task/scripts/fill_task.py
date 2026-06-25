"""
fill_task.py
============
Inserta las respuestas de la Tarea 2 directamente en el .docx original,
SIN modificar ni mover ningún párrafo existente.

Estrategia:
  - Localiza cada párrafo-pregunta por su índice y un fragmento de texto clave.
  - Inserta un bloque de respuesta DESPUÉS del párrafo correspondiente.
  - Copia el estilo de fuente (Times New Roman, 11pt) del documento fuente.
  - Guarda como archivo nuevo para no destruir el original.

Dependencias: python-docx  (ya disponible en el entorno uv del nodo)
Uso:
    python3 scripts/fill_task.py
"""

from __future__ import annotations

import copy
import re
import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Rutas ─────────────────────────────────────────────────────────────────────
VAULT_DIR = Path(__file__).resolve().parents[1]           # docs/vaults/u2-aa-02-poverty-task
SRC_DOCX  = VAULT_DIR / "Tarea 2_U2_DE (1).docx"
OUT_DOCX  = VAULT_DIR / "assets/Tarea2_U2_DE_Respondida.docx"

# ── Respuestas ─────────────────────────────────────────────────────────────────
# Cada entrada: (fragmento_clave_del_párrafo_pregunta, texto_respuesta)
# El fragmento debe identificar ÚNICAMENTE ese párrafo.
RESPUESTAS: list[tuple[str, str]] = [

    # ── PREGUNTA 1 ─────────────────────────────────────────────────────────────
    (
        "Con el boletín del INEC, indique la tasa de pobreza por ingresos",
        (
            "Según el boletín ENEMDU de diciembre 2025 (INEC, 2025), la tasa de pobreza "
            "por ingresos a nivel nacional fue del 17,5 %, mientras que la pobreza "
            "multidimensional alcanzó el 37,8 %. La segunda casi duplica a la primera "
            "porque el enfoque multidimensional trasciende el umbral monetario y captura "
            "privaciones simultáneas en educación, salud, trabajo y hábitat (Tezanos et al., "
            "2013). La medición por ingresos solo determina si el hogar puede adquirir una "
            "canasta básica, ignorando si sus integrantes acceden efectivamente a servicios "
            "públicos, vivienda digna o seguridad laboral. Así, un hogar puede superar la "
            "línea monetaria y, al mismo tiempo, carecer de saneamiento o tener niños fuera "
            "del sistema escolar: privaciones que el enfoque multidimensional registra y "
            "que el monetario no ve (Sen, 1999)."
        ),
    ),

    # ── PREGUNTA 2 ─────────────────────────────────────────────────────────────
    (
        "El Índice de Pobreza Multidimensional del Ecuador asigna un peso de 12,5 %",
        (
            "Incluir el ingreso como componente de una medida declaradamente multidimensional "
            "genera un problema de circularidad endógena: se utiliza el mismo criterio que se "
            "pretende superar como variable explicativa dentro del propio índice. Esta "
            "redundancia viola el principio de independencia conceptual entre dimensiones que "
            "fundamenta el método Alkire-Foster (Alkire & Foster, 2011). Al asignar el 12,5 % "
            "a la pobreza extrema por ingresos, el IPM ecuatoriano hibrida artificialmente dos "
            "paradigmas —el welfarista monetario y el de capacidades— sin resolver la tensión "
            "epistemológica entre ambos. El resultado práctico es que una mejora de ingresos "
            "reduce simultáneamente la privación monetaria y la puntuación en la dimensión "
            "salud-agua-alimentación, inflando la reducción aparente de pobreza multidimensional "
            "más allá de lo que los cambios reales en acceso a servicios justificarían "
            "(Tezanos et al., 2013)."
        ),
    ),

    # ── PREGUNTA 3 ─────────────────────────────────────────────────────────────
    (
        "Entre diciembre de 2024 y diciembre de 2025 la pobreza urbana por ingresos",
        (
            "Una diferencia es estadísticamente significativa cuando la probabilidad de que "
            "sea producto del azar —dado un tamaño de muestra y una varianza determinados— "
            "es suficientemente baja como para rechazar la hipótesis nula de ausencia de "
            "cambio; convencionalmente se usa un nivel de confianza del 95 % (INEC, 2025). "
            "La reducción urbana (7,1 puntos porcentuales) es significativa porque las áreas "
            "urbanas concentran la mayor parte de la muestra ENEMDU, lo que reduce el error "
            "estándar y dota a la estimación de mayor precisión (Tezanos et al., 2013). La "
            "caída rural (5,7 puntos) no alcanza significancia estadística a pesar de ser "
            "numéricamente sustancial: la muestra rural es comparativamente pequeña y la "
            "varianza en los ingresos rurales es elevada, generando intervalos de confianza "
            "amplios que se solapan entre ambos periodos. En términos sustantivos, esto no "
            "implica que la pobreza rural no haya cambiado, sino que la encuesta carece de "
            "la potencia estadística necesaria para confirmar ese cambio con certeza suficiente."
        ),
    ),

    # ── PREGUNTA 4a — Enfoque económico ────────────────────────────────────────
    (
        "Enfoque económico. Calcule el ingreso por persona del hogar",
        (
            "El hogar analizado (cuatro integrantes) percibe un ingreso mensual total de "
            "USD 320. El ingreso per cápita resultante es USD 80,00 (320 ÷ 4), cifra inferior "
            "a la línea de pobreza de diciembre de 2025 (USD 92,40). Por tanto, el hogar "
            "es pobre por ingresos según el criterio monetario del INEC (2025)."
        ),
    ),

    # ── PREGUNTA 4b — Enfoque multidimensional (NBI) ───────────────────────────
    (
        "Enfoque multidimensional. Con los cinco componentes de la pobreza por Necesidades",
        (
            "Aplicando los cinco componentes del NBI (INEC, 2025): (1) Calidad de la "
            "vivienda: sin carencia (paredes y techo en buen estado). (2) Hacinamiento: "
            "sin carencia (cuatro personas en tres ambientes). (3) Servicios básicos: "
            "con carencia (agua por tanquero, sin alcantarillado). (4) Acceso a educación: "
            "sin carencia (menores matriculados). (5) Capacidad económica: con carencia "
            "(ingreso per cápita por debajo de la línea de pobreza y jefe de hogar sin "
            "educación secundaria completa). El hogar presenta carencia en dos de los cinco "
            "componentes; dado que basta con fallar en uno para ser clasificado como pobre "
            "por NBI, el hogar también es pobre por necesidades básicas insatisfechas."
        ),
    ),

    # ── PREGUNTA 4c — ¿Coinciden los diagnósticos? ─────────────────────────────
    (
        "¿Coinciden los dos diagnósticos?",
        (
            "Ambos diagnósticos coinciden: el hogar es pobre tanto por ingresos como por NBI. "
            "Sin embargo, los canales de privación que cada enfoque revela son distintos. El "
            "enfoque monetario capta la insuficiencia inmediata de recursos para cubrir la "
            "canasta básica, mientras que el NBI evidencia la carencia estructural de "
            "saneamiento y la baja capacidad económica del jefe de hogar. Esta coincidencia "
            "sugiere una pobreza de doble capa: transitoria (ingreso) y estructural "
            "(infraestructura), lo que en palabras de Sen (1999) equivale a una privación "
            "acumulada de capacidades que no podría ser detectada con un solo lente de análisis."
        ),
    ),

    # ── PREGUNTA 5 — Decisión de política ──────────────────────────────────────
    (
        "Usted asesora un programa social que debe decidir si ese hogar recibe una ayuda",
        (
            "Para asignar la ayuda priorizaría el enfoque multidimensional por NBI, pues "
            "identifica privaciones estructurales que persisten incluso cuando el ingreso "
            "fluctúa. El acceso a saneamiento y la baja escolaridad del jefe de hogar "
            "constituyen trampas de pobreza de largo plazo que las transferencias monetarias "
            "por sí solas no resuelven (Alkire & Foster, 2011). El enfoque de capacidades "
            "de Sen (1999) refuerza esta postura: el bienestar radica en la capacidad efectiva "
            "de convertir recursos en funcionamientos valiosos, no en la mera posesión de "
            "ingresos. No obstante, descartar el enfoque monetario implicaría perder la señal "
            "más inmediata del consumo. Los programas de transferencia condicionada necesitan "
            "la línea de ingresos para calibrar el monto del beneficio y verificar la "
            "graduación de los beneficiarios; ignorarla puede generar sobreinclusion y "
            "menor eficiencia fiscal (Tezanos et al., 2013)."
        ),
    ),

    # ── PREGUNTA 6 — Sen vs. Nussbaum ──────────────────────────────────────────
    (
        "Sen sostiene que la lista de capacidades básicas debe decidirla cada sociedad",
        (
            "Me inclino por la postura de Nussbaum (2011), quien defiende una lista universal "
            "de capacidades centrales como umbral normativo mínimo que toda sociedad debe "
            "garantizar, con independencia de sus preferencias culturales. Si se deja a cada "
            "comunidad decidir democráticamente qué capacidades son básicas, existe el riesgo "
            "de que estructuras de poder arraigadas —patriarcado, clientelismo— moldeen esa "
            "lista de modo que perpetúen privaciones históricas. La democracia procedimental "
            "no garantiza, por sí sola, resultados sustantivamente justos cuando los grupos "
            "más vulnerables carecen de voz efectiva en el proceso deliberativo. Sin embargo, "
            "reconozco la objeción de Sen (1999): fijar una lista universal puede colisionar "
            "con la agencia y la autodeterminación de las sociedades. Una síntesis operativa "
            "viable consiste en adoptar la lista de Nussbaum como piso normativo no negociable "
            "y dejar que cada sociedad amplíe capacidades adicionales mediante deliberación "
            "democrática."
        ),
    ),

    # ── REFERENCIAS APA ────────────────────────────────────────────────────────
    (
        "Liste, en formato APA, todas las fuentes que citó.",
        (
            "Alkire, S., & Foster, J. (2011). Counting and multidimensional poverty "
            "measurement. Journal of Public Economics, 95(7–8), 476–487. "
            "https://doi.org/10.1016/j.jpubeco.2010.11.006\n\n"
            "Instituto Nacional de Estadística y Censos [INEC]. (2025). Indicadores de "
            "Pobreza y Desigualdad: ENEMDU diciembre 2025. INEC. "
            "https://www.ecuadorencifras.gob.ec/pobreza-y-desigualdad/\n\n"
            "Nussbaum, M. C. (2011). Creating capabilities: The human development approach. "
            "Harvard University Press. https://doi.org/10.4159/harvard.9780674061200\n\n"
            "Sen, A. (1999). Development as freedom. Oxford University Press.\n\n"
            "Tezanos Vázquez, S., Quiñones Montellano, A., & Gutiérrez Sobrao, D. (2013). "
            "Metodologías y herramientas para la medición de la pobreza y la desigualdad. "
            "Cátedra de Cooperación Internacional y con Iberoamérica / Universidad de Cantabria."
        ),
    ),
]


# ── Utilidades XML ─────────────────────────────────────────────────────────────

def _clone_paragraph_format(src_para, dst_para) -> None:
    """Copia el formato de párrafo (pPr) del párrafo fuente al destino."""
    src_pPr = src_para._p.find(qn("w:pPr"))
    if src_pPr is not None:
        dst_pPr = dst_para._p.find(qn("w:pPr"))
        if dst_pPr is None:
            dst_pPr = OxmlElement("w:pPr")
            dst_para._p.insert(0, dst_pPr)
        # Copiar estilo de párrafo si existe
        src_style = src_pPr.find(qn("w:pStyle"))
        if src_style is not None:
            dst_style = dst_pPr.find(qn("w:pStyle"))
            if dst_style is None:
                dst_style = OxmlElement("w:pStyle")
                dst_pPr.insert(0, dst_style)
            dst_style.set(qn("w:val"), src_style.get(qn("w:val"), "Normal"))


def _make_run(para, text: str, bold: bool = False, font_name: str = "Times New Roman", font_size_pt: int = 11):
    """Agrega un run con texto y formato al párrafo dado."""
    run = para.add_run(text)
    run.bold = bold
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    # Forzar fuente también en el XML (para compatibilidad con Word en español)
    r_pr = run._r.get_or_add_rPr()
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:cs"), font_name)
    # Insertar al inicio del rPr si no existe ya
    existing = r_pr.find(qn("w:rFonts"))
    if existing is not None:
        r_pr.remove(existing)
    r_pr.insert(0, r_fonts)
    return run


def _insert_answer_after(doc: Document, after_idx: int, answer_text: str) -> int:
    """
    Inserta uno o más párrafos de respuesta inmediatamente DESPUÉS del párrafo
    en la posición `after_idx` dentro del body del documento.
    Devuelve el número de párrafos insertados.
    """
    body = doc.element.body
    all_paras = body.findall(qn("w:p"))

    if after_idx >= len(all_paras):
        raise IndexError(f"Índice {after_idx} fuera de rango ({len(all_paras)} párrafos).")

    anchor = all_paras[after_idx]
    ref_para = doc.paragraphs[after_idx]  # para copiar formato

    blocks = answer_text.split("\n\n")
    inserted = 0

    # Insertar en orden inverso para mantener posición correcta tras el anchor
    for block in reversed(blocks):
        block = block.strip()
        if not block:
            continue
        new_p = OxmlElement("w:p")
        anchor.addnext(new_p)

        # Crear un párrafo python-docx wrapeando el elemento XML creado
        from docx.text.paragraph import Paragraph
        new_para = Paragraph(new_p, doc)

        # Aplicar formato de párrafo Normal
        _clone_paragraph_format(ref_para, new_para)
        new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _make_run(new_para, block)
        inserted += 1

    return inserted


# ── Main ───────────────────────────────────────────────────────────────────────

def main() -> None:
    if not SRC_DOCX.exists():
        raise FileNotFoundError(f"No se encontró el archivo fuente: {SRC_DOCX}")

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SRC_DOCX, OUT_DOCX)
    print(f"✓ Copia de trabajo creada: {OUT_DOCX}")

    doc = Document(OUT_DOCX)

    # Mapa de fragmento → índice de párrafo (buscamos en tiempo de ejecución)
    def find_para_idx(fragment: str) -> int:
        for i, p in enumerate(doc.paragraphs):
            if fragment.lower() in p.text.lower():
                return i
        raise ValueError(f"Párrafo no encontrado con fragmento: '{fragment[:60]}'")

    # Procesar en orden INVERSO para que los índices no se desplacen
    pairs = []
    for fragment, answer in RESPUESTAS:
        idx = find_para_idx(fragment)
        pairs.append((idx, answer))
        print(f"  → Pregunta localizada en P{idx:03d}: '{fragment[:55]}...'")

    pairs.sort(key=lambda x: x[0], reverse=True)

    for idx, answer in pairs:
        n = _insert_answer_after(doc, idx, answer)
        print(f"  ✓ {n} párrafo(s) insertado(s) después de P{idx:03d}")

    doc.save(OUT_DOCX)
    print(f"\n✅ Documento guardado: {OUT_DOCX}")


if __name__ == "__main__":
    main()
